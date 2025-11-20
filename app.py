import streamlit as st
import speech_recognition as sr
import librosa
import soundfile as sf
from streamlit_mic_recorder import mic_recorder
import io
import os
from docx import Document # Thư viện cho file .docx
import time

# =========================================================================
# KHỞI TẠO VÀ CẤU HÌNH BAN ĐẦU
# =========================================================================
# Khởi tạo đối tượng nhận dạng giọng nói
r = sr.Recognizer()

# Khởi tạo Session State (Quan trọng để giữ trạng thái giữa các lần chạy lại)
if 'audio_buffer' not in st.session_state:
    st.session_state.audio_buffer = None
if 'last_transcription_text' not in st.session_state:
    st.session_state.last_transcription_text = ""
if 'last_audio_data' not in st.session_state:
    st.session_state.last_audio_data = None

# =========================================================================
# HÀM CHUYỂN ĐỔI FILE ÂM THANH (Dùng cho cả File Upload và Ghi âm)
# =========================================================================
def transcribe_audio_from_file_path(file_path):
    """Sử dụng SpeechRecognition để chuyển đổi file WAV thành văn bản."""
    r = sr.Recognizer()
    try:
        with sr.AudioFile(file_path) as source:
            audio = r.record(source) 
        
        text = r.recognize_google(audio, language="vi-VN")
        return text
    except sr.UnknownValueError:
        return "Không thể nhận dạng giọng nói từ tệp âm thanh này."
    except sr.RequestError as e:
        return f"Lỗi kết nối hoặc API: {e}"
    except Exception as e:
        return f"Lỗi xử lý tệp: {e}"

# =========================================================================
# HÀM XỬ LÝ FILE ĐÃ TẢI LÊN (Chuyển đổi sang WAV)
# =========================================================================
def process_uploaded_file(uploaded_file):
    """Xử lý các loại file âm thanh (MP3, WAV,...) thành WAV và chuyển đổi."""
    temp_input_path = "temp_input_audio" + os.path.splitext(uploaded_file.name)[1]
    temp_wav_path = "temp_converted_audio.wav"
    
    st.session_state.last_transcription_text = "" # Reset kết quả cũ

    try:
        with open(temp_input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.info("Đang xử lý và chuyển đổi định dạng âm thanh...")

        # Đọc file bằng librosa
        y, sr_librosa = librosa.load(temp_input_path, sr=None) 
        # Ghi lại thành WAV bằng soundfile
        sf.write(temp_wav_path, y, sr_librosa)
        
        # Chuyển đổi văn bản
        st.info("Đang nhận dạng giọng nói...")
        result_text = transcribe_audio_from_file_path(temp_wav_path)
        
        st.session_state.last_transcription_text = result_text

    except Exception as e:
        st.session_state.last_transcription_text = f"Lỗi xử lý tệp: {e}"
    finally:
        if os.path.exists(temp_input_path):
            os.remove(temp_input_path)
        if os.path.exists(temp_wav_path):
            os.remove(temp_wav_path)

# =========================================================================
# HÀM TẠO VÀ TẢI XUỐNG FILE DOCX
# =========================================================================
def create_docx(text, filename="transcribed_document.docx"):
    """Tạo một file DOCX từ văn bản đã chuyển đổi."""
    document = Document()
    document.add_heading('Văn bản đã chuyển đổi', 0)
    document.add_paragraph(text)

    # Lưu document vào một đối tượng BytesIO
    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io.read(), filename

# =========================================================================
# GIAO DIỆN STREAMLIT
# =========================================================================
def main():
    # Đặt tiêu đề cho ứng dụng
    st.title("🎤 ỨNG DỤNG CHUYỂN ÂM THANH THÀNH VĂN BẢN Ver1.0")
    st.markdown("----------------*************----------------") 
    # Chọn phương thức nhập liệu
    method = st.radio(
    "Chọn phương thức nhập liệu:",
    ('Tải lên File Âm thanh', 'Ghi âm trực tiếp từ Micro')
    )
    ### PHƯƠNG THỨC 1: Tải lên File Âm thanh
    if method == 'Tải lên File Âm thanh':
        uploaded_file = st.file_uploader(
        "Tải lên tệp âm thanh (ví dụ: .wav, .mp3):",
        type=['wav', 'mp3']
        )
        if uploaded_file is not None:
            # Nút chuyển đổi được đặt ở đây
            if st.button('🚀 Chuyển đổi File thành Văn bản'):
                with st.spinner('Đang tải và xử lý file...'):
                    process_uploaded_file(uploaded_file)
    ### PHƯƠNG THỨC 2: Ghi âm trực tiếp từ Micro       
    elif method == 'Ghi âm trực tiếp từ Micro':
        st.subheader("🎙️ Ghi Âm Trực Tiếp")
        st.caption("Sử dụng micro của trình duyệt (thay thế cho PyAudio).")

        # Widget ghi âm
        audio_data = mic_recorder(
            start_prompt="Bắt đầu Ghi Âm",
            stop_prompt="Dừng Ghi Âm",
            key='mic_recorder',
            format="wav" 
        )

        if audio_data:
            # Lưu bytes vào state và hiển thị trình phát
            st.session_state.audio_buffer = audio_data['bytes']
            st.session_state.last_audio_data = audio_data['bytes']
            st.audio(st.session_state.audio_buffer, format='audio/wav')
            
            # Nút TẢI XUỐNG File Âm thanh
            st.download_button(
               label="⬇️ Tải xuống File Âm thanh (.wav)",
               data=st.session_state.last_audio_data,
                file_name="ghi_am_mic.wav",
                mime="audio/wav"
            )
            
        # Logic chuyển đổi văn bản (Chỉ chạy khi có dữ liệu và người dùng bấm nút)
        if st.session_state.audio_buffer is not None:
            if st.button('✅ Chuyển đổi Giọng nói'):
               temp_wav_path = "mic_recording_temp.wav"
                      
                try:
                    # Ghi bytes ra file tạm thời
                    with open(temp_wav_path, "wb") as f:
                        f.write(st.session_state.audio_buffer)

                    # Chuyển đổi văn bản
                    with st.spinner('Đang nhận dạng giọng nói...'):
                        result_text = transcribe_audio_from_file_path(temp_wav_path)

                    st.session_state.last_transcription_text = result_text
                    
                except Exception as e:
                    st.session_state.last_transcription_text = f"Lỗi xử lý: {e}"
                finally:
                    if os.path.exists(temp_wav_path):
                        os.remove(temp_wav_path)
                    # Giữ lại audio_buffer để người dùng có thể tải xuống sau khi chuyển đổi
                # --- Kết thúc Logic Xử lý Ghi âm --- 

                
# =========================================================================
# HIỂN THỊ KẾT QUẢ VÀ TÙY CHỌN TẢI XUỐNG (Chung cho cả 2 phương thức)
# =========================================================================
if st.session_state.last_transcription_text:
    st.markdown("---")
    st.subheader("✅ Văn bản đã chuyển đổi:")

    # Hiển thị văn bản trong textarea
    st.text_area("Kết quả:", st.session_state.last_transcription_text, height=250)
    
    # Chỉ hiển thị nút tải xuống nếu văn bản không phải là lỗi
    if "Không thể" not in st.session_state.last_transcription_text and "Lỗi" not in st.session_state.last_transcription_text:
        
        col1, col2 = st.columns(2)
    
        # Nút tải xuống file DOCX
        docx_bytes, docx_filename = create_docx(st.session_state.last_transcription_text)
        col1.download_button(
            label="💾 Tải xuống MS Word (.docx)",
            data=docx_bytes,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Tùy chọn tải xuống file TXT
        col2.download_button(
            label="📝 Tải xuống Văn bản thuần (.txt)",
            data=st.session_state.last_transcription_text.encode('utf-8'),
            file_name="transcribed_text.txt",
            mime="text/plain"
        )

if __name__ == "__main__":
  main()
  
